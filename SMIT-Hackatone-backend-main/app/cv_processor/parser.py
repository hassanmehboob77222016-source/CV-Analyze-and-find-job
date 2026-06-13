import io

import fitz
from docx import Document


def _extract_pdf_text(file_bytes: bytes) -> str:
    text = ""
    with fitz.open(stream=file_bytes, filetype="pdf") as pdf:
        for page in pdf:
            text += page.get_text()
    return text


def _extract_pdf_text_blocks(file_bytes: bytes) -> str:
    text = ""
    with fitz.open(stream=file_bytes, filetype="pdf") as pdf:
        for page in pdf:
            blocks = page.get_text("blocks")
            for block in blocks:
                if len(block) > 4 and isinstance(block[4], str):
                    text += block[4] + "\n"
    return text


def _ocr_pdf(file_bytes: bytes) -> str:
    text = ""

    try:
        import pytesseract
        from PIL import Image
        try:
            with fitz.open(stream=file_bytes, filetype="pdf") as pdf:
                for page in pdf:
                    pix = page.get_pixmap(dpi=300)
                    mode = "RGBA" if pix.alpha else ("RGB" if pix.n == 3 else "L")
                    image = Image.frombytes(mode, (pix.width, pix.height), pix.samples)
                    if image.mode != "RGB":
                        image = image.convert("RGB")
                    text += pytesseract.image_to_string(image, lang="eng")
            if len(text.strip()) >= 50:
                return text
        except (pytesseract.TesseractNotFoundError, Exception):
            text = ""
    except ImportError:
        text = ""

    if len(text.strip()) >= 50:
        return text

    try:
        import easyocr
        import numpy as np
        from PIL import Image
    except ImportError:
        return text

    try:
        reader = easyocr.Reader(["en"], gpu=False)
        with fitz.open(stream=file_bytes, filetype="pdf") as pdf:
            for page in pdf:
                pix = page.get_pixmap(dpi=300)
                mode = "RGBA" if pix.alpha else ("RGB" if pix.n == 3 else "L")
                image = Image.frombytes(mode, (pix.width, pix.height), pix.samples)
                if image.mode != "RGB":
                    image = image.convert("RGB")
                image_np = np.array(image)
                page_text = "\n".join(reader.readtext(image_np, detail=0) or [])
                text += page_text + "\n"
        return text
    except Exception:
        return text


def extract_all_docx_text(file_bytes: bytes) -> str:
    doc = Document(io.BytesIO(file_bytes))
    text_parts = []

    # 1. Regular paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text.strip())

    # 2. Tables (many CVs use tables for layout)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.text.strip():
                        text_parts.append(para.text.strip())

    # 3. Headers and footers
    for section in doc.sections:
        for header_para in section.header.paragraphs:
            if header_para.text.strip():
                text_parts.append(header_para.text.strip())
        for footer_para in section.footer.paragraphs:
            if footer_para.text.strip():
                text_parts.append(footer_para.text.strip())

    # 4. Text boxes (stored in XML as txbxContent)
    for shape in doc.element.body.iter():
        if shape.tag.endswith("}txbxContent"):
            for child in shape.iter():
                if child.tag.endswith("}t") and child.text and child.text.strip():
                    text_parts.append(child.text.strip())

    # Remove duplicates while keeping order
    seen = set()
    unique_parts = []
    for part in text_parts:
        if part not in seen:
            seen.add(part)
            unique_parts.append(part)

    return "\n".join(unique_parts)


def parse_cv(file_bytes: bytes, filename: str) -> str:
    ext = filename.lower().split(".")[-1]

    if ext == "pdf":
        try:
            text = _extract_pdf_text(file_bytes)
        except Exception as exc:
            raise ValueError(f"PDF read error: {str(exc)}") from exc

        if len(text.strip()) < 50:
            try:
                text = _extract_pdf_text_blocks(file_bytes)
            except Exception:
                text = ""

        if len(text.strip()) < 50:
            try:
                ocr_text = _ocr_pdf(file_bytes)
            except Exception:
                ocr_text = ""

            if ocr_text.strip():
                text = ocr_text

        if not text.strip():
            raise ValueError(
                "Could not extract text from this PDF. "
                "Your PDF appears to be scanned/image-based. "
                "Install Tesseract OCR and the Python packages `pytesseract` and `Pillow`, "
                "or install `easyocr` and `opencv-python-headless` for a Python-only OCR fallback. "
                "Alternatively use a searchable PDF exported from Word/Google Docs."
            )

        return text.strip()


        return text.strip()

    if ext in ["docx", "doc"]:
        try:
            text = extract_all_docx_text(file_bytes)
            if len(text.strip()) < 50:
                raise ValueError(
                    "Word document appears to be empty or uses unsupported formatting. "
                    "Please make sure your CV has actual typed text (not just images)."
                )
            return text.strip()
        except Exception as exc:
            raise ValueError(f"Word document read error: {str(exc)}") from exc

    raise ValueError(f"Unsupported file type: {ext}. Use PDF or DOCX only.")
